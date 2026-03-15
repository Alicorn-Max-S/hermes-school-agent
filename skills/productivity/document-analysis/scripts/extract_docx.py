#!/usr/bin/env python3
"""Extract text from DOCX files using python-docx (FREE, local).

Usage:
    python3 extract_docx.py document.docx
    python3 extract_docx.py document.docx --format markdown
    python3 extract_docx.py document.docx --format text
"""

import json
import sys
from pathlib import Path


def extract(input_path, fmt="markdown"):
    """Extract text from a DOCX file."""
    path = Path(input_path)
    if not path.exists():
        return {"success": False, "error": f"File not found: {input_path}"}

    try:
        from docx import Document
    except ImportError:
        return {
            "success": False,
            "error": "python-docx not installed. Run: pip install python-docx",
        }

    try:
        doc = Document(str(path))
    except Exception as e:
        return {"success": False, "error": f"Failed to open DOCX: {e}"}

    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        if fmt == "markdown":
            style = (para.style.name or "").lower()
            if "heading 1" in style:
                lines.append(f"# {text}")
            elif "heading 2" in style:
                lines.append(f"## {text}")
            elif "heading 3" in style:
                lines.append(f"### {text}")
            elif "heading 4" in style:
                lines.append(f"#### {text}")
            elif "list" in style:
                lines.append(f"- {text}")
            else:
                lines.append(text)
        else:
            lines.append(text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        lines.append("")
        if fmt == "markdown":
            lines.append(f"**Table {i + 1}:**")
        else:
            lines.append(f"Table {i + 1}:")

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if fmt == "markdown":
                lines.append("| " + " | ".join(cells) + " |")
                if row_idx == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            else:
                lines.append("\t".join(cells))

    content = "\n".join(lines).strip()

    return {
        "success": True,
        "content": content,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "format": fmt,
    }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: extract_docx.py <file> [--format markdown|text]"}))
        sys.exit(1)

    fmt = "markdown"
    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            fmt = sys.argv[idx + 1]

    result = extract(sys.argv[1], fmt)
    print(json.dumps(result, indent=2, ensure_ascii=False))
